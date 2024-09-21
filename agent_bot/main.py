import os
from dotenv import load_dotenv
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, MessageHandler, CallbackQueryHandler, filters, ContextTypes
from telegram.constants import ChatAction
from qdrant_client import models, QdrantClient
from sentence_transformers import SentenceTransformer
import google.generativeai as genai
import sqlite3
import asyncio
import json
from docx import Document
from docx2pdf import convert
from PyPDF2 import PdfReader, PdfWriter

# Load environment variables
load_dotenv()

# Set up the Gemini API
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# Initialize the sentence transformer model
encoder = SentenceTransformer("all-MiniLM-L6-v2")

# Set up Qdrant client
qdrant_storage_path = "./qdrant_storage"
client = QdrantClient(path=qdrant_storage_path)
collection_name = "real_estate_data"

# Set up agent and property databases
AGENT_DB_PATH = 'agent_data.db'
PROPERTY_DB_PATH = 'property_data.db'

# Set up chat history and user states
CHAT_HISTORIES = {}
USER_STATES = {}


def setup_databases():
    conn = sqlite3.connect(AGENT_DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS agents (
        agent_id INTEGER PRIMARY KEY,
        name TEXT,
        agency TEXT,
        last_activity TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )
    ''')
    conn.commit()
    conn.close()

    conn = sqlite3.connect(PROPERTY_DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS properties (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        address TEXT,
        price REAL,
        bedrooms INTEGER,
        bathrooms INTEGER,
        square_feet REAL,
        description TEXT,
        agent_id INTEGER
    )
    ''')
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS clients (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT,
        email TEXT,
        phone TEXT,
        agent_id INTEGER
    )
    ''')
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS documents (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        type TEXT,
        content TEXT,
        property_id INTEGER,
        client_id INTEGER,
        agent_id INTEGER,
        status TEXT
    )
    ''')
    conn.commit()
    conn.close()


def get_or_create_agent(agent_id, name, agency):
    conn = sqlite3.connect(AGENT_DB_PATH)
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM agents WHERE agent_id = ?', (agent_id,))
    result = cursor.fetchone()
    if result is None:
        cursor.execute(
            'INSERT INTO agents (agent_id, name, agency) VALUES (?, ?, ?)',
            (agent_id, name, agency)
        )
        conn.commit()
    else:
        cursor.execute(
            'UPDATE agents SET last_activity = CURRENT_TIMESTAMP WHERE agent_id = ?',
            (agent_id,)
        )
        conn.commit()
    conn.close()


def add_property(address, price, bedrooms, bathrooms, square_feet, description, agent_id):
    conn = sqlite3.connect(PROPERTY_DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''
    INSERT INTO properties (address, price, bedrooms, bathrooms, square_feet, description, agent_id)
    VALUES (?, ?, ?, ?, ?, ?, ?)
    ''', (address, price, bedrooms, bathrooms, square_feet, description, agent_id))
    property_id = cursor.lastrowid
    conn.commit()
    conn.close()
    return property_id


def get_property(property_id):
    conn = sqlite3.connect(PROPERTY_DB_PATH)
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM properties WHERE id = ?', (property_id,))
    result = cursor.fetchone()
    conn.close()
    return result


def add_client(name, email, phone, agent_id):
    conn = sqlite3.connect(PROPERTY_DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''
    INSERT INTO clients (name, email, phone, agent_id)
    VALUES (?, ?, ?, ?)
    ''', (name, email, phone, agent_id))
    client_id = cursor.lastrowid
    conn.commit()
    conn.close()
    return client_id


def get_client(client_id):
    conn = sqlite3.connect(PROPERTY_DB_PATH)
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM clients WHERE id = ?', (client_id,))
    result = cursor.fetchone()
    conn.close()
    return result


def create_document(doc_type, content, property_id, client_id, agent_id):
    conn = sqlite3.connect(PROPERTY_DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''
    INSERT INTO documents (type, content, property_id, client_id, agent_id, status)
    VALUES (?, ?, ?, ?, ?, ?)
    ''', (doc_type, content, property_id, client_id, agent_id, 'draft'))
    document_id = cursor.lastrowid
    conn.commit()
    conn.close()
    return document_id


def get_document(document_id):
    conn = sqlite3.connect(PROPERTY_DB_PATH)
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM documents WHERE id = ?', (document_id,))
    result = cursor.fetchone()
    conn.close()
    return result


def update_document_status(document_id, status):
    conn = sqlite3.connect(PROPERTY_DB_PATH)
    cursor = conn.cursor()
    cursor.execute('UPDATE documents SET status = ? WHERE id = ?',
                   (status, document_id))
    conn.commit()
    conn.close()


def generate_document(doc_type, property_data, client_data):
    # This is a placeholder function. In a real-world scenario, you'd use a template engine
    # or a more sophisticated document generation system.
    doc = Document()
    doc.add_heading(f"{doc_type.capitalize()} Agreement", 0)

    doc.add_paragraph(f"Property Address: {property_data['address']}")
    doc.add_paragraph(f"Price: ${property_data['price']}")
    doc.add_paragraph(f"Client: {client_data['name']}")
    doc.add_paragraph(f"Client Email: {client_data['email']}")
    doc.add_paragraph(f"Client Phone: {client_data['phone']}")

    doc.add_paragraph("This is a placeholder for the full agreement text.")

    filename = f"{doc_type.lower().replace(' ', '_')}.docx"
    doc.save(filename)

    # Convert to PDF
    pdf_filename = filename.replace('.docx', '.pdf')
    convert(filename, pdf_filename)

    return pdf_filename


async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user = update.effective_user
    get_or_create_agent(user.id, user.full_name, "Unknown Agency")

    CHAT_HISTORIES[user.id] = []
    USER_STATES[user.id] = {'state': 'idle'}

    keyboard = [
        [InlineKeyboardButton("Add Property", callback_data='add_property')],
        [InlineKeyboardButton("Add Client", callback_data='add_client')],
        [InlineKeyboardButton("Prepare Document",
                              callback_data='prepare_document')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)

    await update.message.reply_text(
        f'Welcome, {user.full_name}! How can I assist you today?',
        reply_markup=reply_markup
    )


async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    user = update.effective_user
    get_or_create_agent(user.id, user.full_name, "Unknown Agency")

    await context.bot.send_chat_action(chat_id=update.effective_chat.id, action=ChatAction.TYPING)

    query = update.message.text
    state = USER_STATES[user.id]['state']

    if state == 'add_property':
        # Handle property addition logic
        property_data = json.loads(query)
        property_id = add_property(
            property_data['address'],
            property_data['price'],
            property_data['bedrooms'],
            property_data['bathrooms'],
            property_data['square_feet'],
            property_data['description'],
            user.id
        )
        USER_STATES[user.id]['state'] = 'idle'
        await update.message.reply_text(f"Property added successfully! Property ID: {property_id}")
    elif state == 'add_client':
        # Handle client addition logic
        client_data = json.loads(query)
        client_id = add_client(
            client_data['name'],
            client_data['email'],
            client_data['phone'],
            user.id
        )
        USER_STATES[user.id]['state'] = 'idle'
        await update.message.reply_text(f"Client added successfully! Client ID: {client_id}")
    elif state == 'prepare_document':
        # Handle document preparation logic
        doc_data = json.loads(query)
        property_data = get_property(doc_data['property_id'])
        client_data = get_client(doc_data['client_id'])

        if property_data and client_data:
            pdf_filename = generate_document(
                doc_data['type'], property_data, client_data)
            document_id = create_document(
                doc_data['type'],
                pdf_filename,
                doc_data['property_id'],
                doc_data['client_id'],
                user.id
            )
            USER_STATES[user.id]['state'] = 'idle'
            await update.message.reply_document(document=open(pdf_filename, 'rb'))
            await update.message.reply_text(f"Document prepared successfully! Document ID: {document_id}")
        else:
            await update.message.reply_text("Error: Property or client not found.")
    else:
        # General query handling
        agent_history = CHAT_HISTORIES.get(user.id, [])
        chat_history = "\n".join(agent_history)

        improved_query = improve_query(query, chat_history)
        response, hits = retrieve_and_generate(improved_query, chat_history)

        agent_history = agent_history[-4:] + \
            [f"Agent: {query}", f"Bot: {response}"]
        CHAT_HISTORIES[user.id] = agent_history

        await update.message.reply_text(response, disable_web_page_preview=True)


async def button_callback(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    query = update.callback_query
    await query.answer()

    if query.data == 'add_property':
        USER_STATES[query.from_user.id]['state'] = 'add_property'
        await query.edit_message_text("Please provide the property details in JSON format:")
    elif query.data == 'add_client':
        USER_STATES[query.from_user.id]['state'] = 'add_client'
        await query.edit_message_text("Please provide the client details in JSON format:")
    elif query.data == 'prepare_document':
        USER_STATES[query.from_user.id]['state'] = 'prepare_document'
        await query.edit_message_text("Please provide the document details in JSON format:")


def main() -> None:
    setup_databases()

    application = Application.builder().token(
        os.getenv("TELEGRAM_BOT_TOKEN")).build()

    application.add_handler(CommandHandler("start", start))
    application.add_handler(MessageHandler(
        filters.TEXT & ~filters.COMMAND, handle_message))
    application.add_handler(CallbackQueryHandler(button_callback))

    application.run_polling(allowed_updates=Update.ALL_TYPES)


if __name__ == '__main__':
    main()
